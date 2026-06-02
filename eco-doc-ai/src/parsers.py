import logging
import os
import platform
import re
import subprocess
import tempfile
import uuid
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

import fitz
import pandas as pd
from docx import Document

from pdf_page_quality import (
    evaluate_page_text_quality,
    page_has_embedded_images,
    postprocess_ocr_text,
    should_attempt_ocr,
    try_complementary_merge,
)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".heic"}

logger = logging.getLogger(__name__)

SCHEMA_NAME = "eco_doc_record"
SCHEMA_VERSION = "2.0"

_docling_converter = None

# общий маркер страниц для pdf (docling) и docx; одна константа для дальнейшей подстройки.
DEFAULT_PAGE_MARKER = "\n\n--- Page {n} ---\n\n"


@dataclass
class PdfExtractionConfig:
    """
    извлечение pdf по страницам; параметры меняйте здесь, без правки внутренностей parse_pdf.

    атрибуты:
        min_native_quality_score: порог качества нативного текста (метаданные страницы).
        max_ocr_pages_per_document: лимит вызовов ocr (none = без лимита).
        ocr_dpi: dpi растра для ocr.
        ocr_score_advantage: ocr должен обогнать натив по score минимум на эту величину.
        complementary_merge: разрешить осторожное слияние натив+ocr (режим mixed).
        page_marker: вставляется перед телом каждой страницы в собранном тексте.
        pdf_engine: docling — разбор с учётом вёрстки (таблицы, порядок чтения);
            pymupdf — прежний пайплайн pymupdf + paddleocr по страницам.
    """

    min_native_quality_score: float = 0.38
    max_ocr_pages_per_document: int | None = None
    ocr_dpi: int = 150
    ocr_score_advantage: float = 0.04
    complementary_merge: bool = True
    page_marker: str = DEFAULT_PAGE_MARKER
    pdf_engine: str = "docling"

# убираем лишние пробелы и переносы
def _clean_whitespace(text: str) -> str:
    text = (text or "").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# функция для декодирования
def _read_text_with_fallbacks(path: str, encodings: tuple[str, ...]) -> str:
    raw = Path(path).read_bytes()
    for enc in encodings:
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1", errors="replace")

# делаем одинаковую базу для всех файлов
def _make_base_result(path: str, file_type: str) -> dict:
    file_path = Path(path)
    return {
        "document_id": str(uuid.uuid5(uuid.NAMESPACE_URL, str(file_path.resolve()))),
        "schema_name": SCHEMA_NAME,
        "schema_version": SCHEMA_VERSION,
        "source_file": file_path.name,
        "file_type": file_type,
        "text": "",
        "paragraphs": [],
        "pages": [],
        "tables": [],
        "metadata": {
            "ocr_used": False,
            "ocr_engine": None,
            "ocr_pages": 0,
        },
    }

def _pdf_page_to_png(page: fitz.Page, dpi: int = 150) -> Path:
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    fd, tmp_path = tempfile.mkstemp(suffix=".png")
    import os

    os.close(fd)
    pix.save(tmp_path)
    return Path(tmp_path)


def _paddle_lines_from_ocr_result(result) -> list[str]:
    lines: list[str] = []
    for block in result or []:
        if not block:
            continue
        for item in block:
            text = item[1][0].strip()
            if text:
                lines.append(text)
    return lines


def _ocr_single_pdf_page(page: fitz.Page, ocr, dpi: int) -> str:
    from ocr import run_ocr_on_image

    tmp = _pdf_page_to_png(page, dpi=dpi)
    try:
        result = run_ocr_on_image(ocr, str(tmp))
        return "\n".join(_paddle_lines_from_ocr_result(result)).strip()
    finally:
        tmp.unlink(missing_ok=True)


def _pick_native_vs_ocr(
    native_raw: str,
    ocr_clean: str,
    qn,
    qo,
    cfg: PdfExtractionConfig,
) -> tuple[str, str]:
    """выбирает лучший итоговый текст страницы и режим extraction_mode."""
    n = (native_raw or "").strip()
    o = (ocr_clean or "").strip()

    if not n and not o:
        return "", "failed"
    if not o:
        return native_raw.strip(), "native" if n else "failed"
    if not n:
        return ocr_clean.strip(), "ocr"

    merged, merge_skip = try_complementary_merge(native_raw, ocr_clean)
    if (
        cfg.complementary_merge
        and merged is not None
        and merge_skip is None
    ):
        qm = evaluate_page_text_quality(merged)
        if qm.score >= max(qn.score, qo.score) - 0.08 and qm.char_len >= min(
            qn.char_len, qo.char_len
        ):
            return merged.strip(), "mixed"

    if qo.score >= qn.score + cfg.ocr_score_advantage and qo.char_len >= 15:
        return ocr_clean.strip(), "ocr"
    if qn.score >= qo.score - 1e-6:
        return native_raw.strip(), "native"
    return ocr_clean.strip(), "ocr"


def _docling_available() -> bool:
    try:
        import docling  # noqa: F401

        return True
    except ImportError:
        return False


def get_docling_converter():
    """ленивый синглтон для всех форматов через docling (pdf, docx, …)."""
    global _docling_converter
    if _docling_converter is None:
        from docling.document_converter import DocumentConverter

        # позже: documentconverter(..., format_options={...}) без протаскивания конфига по вызовам.
        _docling_converter = DocumentConverter()
    return _docling_converter


def _docling_table_as_pipe_rows(tbl: dict) -> list[str]:
    data = tbl.get("data") or {}
    cells = data.get("table_cells") or []
    if not cells:
        return []
    by_row: dict[int, list[tuple[int, str]]] = defaultdict(list)
    for c in cells:
        if not isinstance(c, dict):
            continue
        r = int(c.get("start_row_offset_idx", 0))
        col = int(c.get("start_col_offset_idx", 0))
        txt = (c.get("text") or "").replace("\n", " ").strip()
        by_row.setdefault(r, []).append((col, txt))
    lines: list[str] = []
    for r in sorted(by_row.keys()):
        row_cells = [t for _, t in sorted(by_row[r], key=lambda x: x[0])]
        lines.append(" | ".join(row_cells))
    return lines


def _docling_dict_tables_to_eco(tbls: list) -> list[dict]:
    out: list[dict] = []
    for table_idx, tbl in enumerate(tbls or []):
        if not isinstance(tbl, dict):
            continue
        data = tbl.get("data") or {}
        cells = data.get("table_cells") or []
        if not cells:
            continue
        by_row: dict[int, list[tuple[int, str]]] = defaultdict(list)
        for c in cells:
            if not isinstance(c, dict):
                continue
            r = int(c.get("start_row_offset_idx", 0))
            col = int(c.get("start_col_offset_idx", 0))
            txt = (c.get("text") or "").replace("\n", " ").strip()
            by_row.setdefault(r, []).append((col, txt))
        table_rows: list[dict] = []
        for r in sorted(by_row.keys()):
            row_cells = [t for _, t in sorted(by_row[r], key=lambda x: x[0])]
            table_rows.append({"row_num": r + 1, "cells": row_cells})
        out.append({"table_num": table_idx + 1, "rows": table_rows})
    return out


def _paragraphs_from_docling_texts(data: dict) -> list[str]:
    """блоки текста в порядке чтения; у docx в экспорте часто нет prov — сохраняем порядок списка."""
    blocks: list[tuple[int, float, float, str]] = []
    seq = 0
    for item in data.get("texts") or []:
        if not isinstance(item, dict):
            continue
        t = (item.get("text") or "").strip()
        if not t:
            continue
        provs = item.get("prov") or []
        if provs:
            for prov in provs:
                pno = int(prov.get("page_no", 1))
                bbox = prov.get("bbox") or {}
                top = float(bbox.get("t", 0.0))
                left = float(bbox.get("l", 0.0))
                blocks.append((pno, -top, left, t))
        else:
            seq += 1
            blocks.append((1, -float(seq), 0.0, t))
    blocks.sort(key=lambda x: (x[0], x[1], x[2]))
    return [b[3] for b in blocks]


def _parse_with_docling(path: str, record_file_type: str, page_marker: str) -> dict:
    """разбор через docling с учётом вёрстки (pdf, docx): порядок чтения, таблицы, страницы."""
    out = _make_base_result(path, record_file_type)
    converter = get_docling_converter()
    result = converter.convert(path)
    doc = result.document
    data = doc.export_to_dict()

    page_keys = data.get("pages") or {}
    page_nums = sorted(int(k) for k in page_keys.keys())
    if not page_nums:
        inferred: set[int] = set()
        for item in data.get("texts") or []:
            if not isinstance(item, dict):
                continue
            for pr in item.get("prov") or []:
                inferred.add(int(pr.get("page_no", 1)))
        for tbl in data.get("tables") or []:
            if not isinstance(tbl, dict):
                continue
            for pr in tbl.get("prov") or []:
                inferred.add(int(pr.get("page_no", 1)))
        page_nums = sorted(inferred) or [1]

    by_page_lines: dict[int, list[tuple[float, float, str]]] = defaultdict(list)
    text_seq = 0
    fallback_page = page_nums[0] if page_nums else 1
    for item in data.get("texts") or []:
        if not isinstance(item, dict):
            continue
        t = (item.get("text") or "").strip()
        if not t:
            continue
        provs = item.get("prov") or []
        if provs:
            for prov in provs:
                pno = int(prov.get("page_no", 1))
                bbox = prov.get("bbox") or {}
                top = float(bbox.get("t", 0.0))
                left = float(bbox.get("l", 0.0))
                by_page_lines[pno].append((-top, left, t))
        else:
            text_seq += 1
            by_page_lines[fallback_page].append((-float(text_seq), 0.0, t))

    table_blocks_by_page: dict[int, list[tuple[float, str]]] = defaultdict(list)
    for tbl in data.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        lines = _docling_table_as_pipe_rows(tbl)
        if not lines:
            continue
        block = "\n".join(lines)
        provs = tbl.get("prov") or []
        if provs:
            pno = int(provs[0].get("page_no", 1))
            bbox = provs[0].get("bbox") or {}
            top = float(bbox.get("t", 0.0))
        else:
            pno = page_nums[0]
            top = 0.0
        table_blocks_by_page[pno].append((-top, block))

    pages_out: list[dict] = []
    failed_pages = 0
    for pno in page_nums:
        chunks = sorted(by_page_lines.get(pno, []), key=lambda x: (x[0], x[1]))
        text_chunks = [c[2] for c in chunks]
        tbl_chunks = sorted(table_blocks_by_page.get(pno, []), key=lambda x: x[0])
        table_texts = [b for _, b in tbl_chunks]
        parts: list[str] = []
        if text_chunks:
            parts.append("\n".join(text_chunks))
        if table_texts:
            parts.append("\n\n".join(table_texts))
        final_text = "\n\n".join(parts).strip()

        if not final_text:
            failed_pages += 1
            quality_flag = "failed"
            mode = "failed"
            q_final_score = 0.0
        else:
            q_final = evaluate_page_text_quality(final_text)
            q_final_score = round(q_final.score, 4)
            mode = "docling"
            if q_final.score < 0.32 and q_final.char_len > 40:
                quality_flag = "low"
            else:
                quality_flag = "ok"

        page_entry = {
            "page_num": pno,
            "text": final_text,
            "extraction_mode": mode,
            "native_text_length": 0,
            "final_text_length": len(final_text),
            "native_quality_score": 0.0,
            "final_quality_score": q_final_score,
            "ocr_applied": False,
            "quality_flag": quality_flag,
            "reason_for_ocr": None,
        }
        pages_out.append(page_entry)

    text_parts: list[str] = []
    for p in pages_out:
        text_parts.append(page_marker.format(n=p["page_num"]))
        body = p["text"] or ""
        text_parts.append(_clean_whitespace(body) if body else "")

    assembled = "".join(text_parts).strip()
    out["text"] = _clean_whitespace(assembled)
    out["pages"] = pages_out
    out["tables"] = _docling_dict_tables_to_eco(data.get("tables"))
    out["paragraphs"] = _paragraphs_from_docling_texts(data)

    try:
        import docling as _dl

        dl_ver = getattr(_dl, "__version__", "unknown")
    except Exception:
        dl_ver = "unknown"

    summary = (
        f"docling total_pages={len(page_nums)} ok_pages={len(page_nums) - failed_pages} "
        f"failed_pages={failed_pages} tables={len(out['tables'])}"
    )
    out["metadata"]["ocr_used"] = False
    out["metadata"]["ocr_engine"] = None
    out["metadata"]["ocr_pages"] = 0
    out["metadata"]["total_pages"] = len(page_nums)
    out["metadata"]["native_pages_count"] = 0
    out["metadata"]["ocr_pages_count"] = 0
    out["metadata"]["mixed_pages_count"] = 0
    out["metadata"]["failed_pages_count"] = failed_pages
    out["metadata"]["extraction_summary"] = summary
    if record_file_type == "pdf":
        out["metadata"]["pdf_extraction"] = "docling_v1"
    else:
        out["metadata"]["docx_extraction"] = "docling_v1"
    out["metadata"]["docling_version"] = dl_ver
    out["metadata"]["fallback_reason"] = None
    logger.info(
        "Docling extraction (%s) %s: %s",
        record_file_type,
        Path(path).name,
        summary,
    )
    return out


def _parse_pdf_docling(path: str, cfg: PdfExtractionConfig) -> dict:
    return _parse_with_docling(path, "pdf", cfg.page_marker)


def _parse_pdf_pymupdf(path: str, config: PdfExtractionConfig | None = None) -> dict:
    """
    извлечение pdf по страницам: на каждой странице отдельно натив или ocr.
    совместимо с eco_doc_record: text, pages[], расширенные metadata.
    """
    cfg = config or PdfExtractionConfig()
    out = _make_base_result(path, "pdf")
    doc = fitz.open(path)
    ocr_engine = None
    ocr_available = True
    ocr_invocations = 0
    logged_ocr_unavailable_skip = False

    pages_out: list[dict] = []
    native_mode_pages = 0
    ocr_mode_pages = 0
    mixed_mode_pages = 0
    failed_pages = 0
    fallback_reasons: list[str] = []

    try:
        total_pages = len(doc)
        for i in range(total_pages):
            page = doc[i]
            page_num = i + 1
            native_raw = page.get_text("text").strip()
            qn = evaluate_page_text_quality(native_raw)
            has_images = page_has_embedded_images(page)
            need_ocr, ocr_reason = should_attempt_ocr(native_raw, has_images)

            ocr_applied = False
            reason_for_ocr: str | None = None
            final_text = ""
            mode = "native"
            ocr_clean = ""

            if not need_ocr:
                final_text = native_raw
                mode = "native"
                logger.debug(
                    "PDF %s p.%s: native OK (len=%s score=%.3f)",
                    Path(path).name,
                    page_num,
                    qn.char_len,
                    qn.score,
                )
            else:
                reason_for_ocr = ocr_reason

                budget_ok = True
                if cfg.max_ocr_pages_per_document is not None:
                    if ocr_invocations >= cfg.max_ocr_pages_per_document:
                        budget_ok = False
                        if "ocr_budget_exhausted" not in fallback_reasons:
                            fallback_reasons.append("ocr_budget_exhausted")

                if not budget_ok:
                    final_text = native_raw
                    mode = "native" if native_raw.strip() else "failed"
                    reason_for_ocr = (reason_for_ocr or "ocr_needed") + "|ocr_budget_exhausted"
                    logger.warning(
                        "PDF %s p.%s: OCR skipped (budget); native len=%s",
                        Path(path).name,
                        page_num,
                        qn.char_len,
                    )
                elif not ocr_available:
                    final_text = native_raw
                    mode = "native" if native_raw.strip() else "failed"
                    reason_for_ocr = (reason_for_ocr or "ocr_needed") + "|ocr_unavailable"
                    if not logged_ocr_unavailable_skip:
                        logger.warning(
                            "PDF %s: further pages skip OCR (engine unavailable); "
                            "example page %s native len=%s",
                            Path(path).name,
                            page_num,
                            qn.char_len,
                        )
                        logged_ocr_unavailable_skip = True
                    else:
                        logger.debug(
                            "PDF %s p.%s: OCR skipped (engine unavailable); native len=%s",
                            Path(path).name,
                            page_num,
                            qn.char_len,
                        )
                else:
                    if ocr_engine is None:
                        try:
                            from ocr import get_ocr

                            ocr_engine = get_ocr()
                        except Exception as e:
                            ocr_available = False
                            if "ocr_engine_unavailable" not in fallback_reasons:
                                fallback_reasons.append("ocr_engine_unavailable")
                            logger.error(
                                "PDF %s: OCR engine init failed, native-only fallback: %s",
                                Path(path).name,
                                e,
                            )
                            final_text = native_raw
                            mode = "native" if native_raw.strip() else "failed"
                            reason_for_ocr = (reason_for_ocr or "ocr_needed") + "|ocr_init_failed"
                    if ocr_engine is not None:
                        try:
                            ocr_raw = _ocr_single_pdf_page(page, ocr_engine, cfg.ocr_dpi)
                        except Exception as e:
                            logger.warning(
                                "PDF %s p.%s: OCR error: %s",
                                Path(path).name,
                                page_num,
                                e,
                            )
                            ocr_raw = ""
                        ocr_applied = True
                        ocr_invocations += 1
                        ocr_clean = postprocess_ocr_text(ocr_raw)
                        qo = evaluate_page_text_quality(ocr_clean)
                        final_text, mode = _pick_native_vs_ocr(
                            native_raw, ocr_clean, qn, qo, cfg
                        )
                        logger.debug(
                            "PDF %s p.%s: OCR applied; mode=%s native_score=%.3f ocr_score=%.3f",
                            Path(path).name,
                            page_num,
                            mode,
                            qn.score,
                            qo.score,
                        )

            final_text = (final_text or "").strip()
            if not final_text:
                mode = "failed"

            if mode == "native":
                native_mode_pages += 1
            elif mode == "ocr":
                ocr_mode_pages += 1
            elif mode == "mixed":
                mixed_mode_pages += 1
            else:
                failed_pages += 1

            q_final = evaluate_page_text_quality(final_text)
            if mode == "failed" or not final_text:
                quality_flag = "failed"
            elif mode == "native" and qn.score < cfg.min_native_quality_score:
                quality_flag = "low"
            elif q_final.score < 0.32 and len((final_text or "").strip()) > 40:
                quality_flag = "low"
            else:
                quality_flag = "ok"

            page_entry = {
                "page_num": page_num,
                "text": final_text,
                "extraction_mode": mode,
                "native_text_length": len(native_raw),
                "final_text_length": len(final_text),
                "native_quality_score": round(qn.score, 4),
                "final_quality_score": round(q_final.score, 4),
                "ocr_applied": ocr_applied,
                "quality_flag": quality_flag,
                "reason_for_ocr": reason_for_ocr if need_ocr else None,
            }
            pages_out.append(page_entry)

        text_parts: list[str] = []
        for p in pages_out:
            text_parts.append(cfg.page_marker.format(n=p["page_num"]))
            body = p["text"] or ""
            text_parts.append(_clean_whitespace(body) if body else "")

        assembled = "".join(text_parts).strip()
        out["text"] = _clean_whitespace(assembled)
        out["pages"] = pages_out

        any_ocr = ocr_invocations > 0
        summary = (
            f"page_level total={total_pages} native={native_mode_pages} "
            f"ocr={ocr_mode_pages} mixed={mixed_mode_pages} failed={failed_pages} "
            f"ocr_calls={ocr_invocations}"
        )
        out["metadata"]["ocr_used"] = any_ocr
        out["metadata"]["ocr_engine"] = "paddleocr" if any_ocr else None
        out["metadata"]["ocr_pages"] = ocr_invocations
        out["metadata"]["total_pages"] = total_pages
        out["metadata"]["native_pages_count"] = native_mode_pages
        out["metadata"]["ocr_pages_count"] = ocr_mode_pages
        out["metadata"]["mixed_pages_count"] = mixed_mode_pages
        out["metadata"]["failed_pages_count"] = failed_pages
        out["metadata"]["extraction_summary"] = summary
        out["metadata"]["pdf_extraction"] = "page_level_v1"
        out["metadata"]["fallback_reason"] = (
            ";".join(fallback_reasons) if fallback_reasons else None
        )
        logger.info("PDF extraction %s: %s", Path(path).name, summary)
    finally:
        doc.close()

    return out


def parse_pdf(path: str, config: PdfExtractionConfig | None = None) -> dict:
    cfg = config or PdfExtractionConfig()
    env_engine = (os.environ.get("ECO_DOC_PDF_ENGINE") or "").strip().lower()
    engine = env_engine if env_engine in ("docling", "pymupdf") else cfg.pdf_engine

    if engine == "docling":
        if _docling_available():
            try:
                return _parse_pdf_docling(path, cfg)
            except Exception as e:
                logger.warning(
                    "Docling PDF parse failed (%s), falling back to PyMuPDF: %s",
                    Path(path).name,
                    e,
                )
                return _parse_pdf_pymupdf(path, cfg)
        logger.warning(
            "pdf_engine=docling but docling is not installed; using PyMuPDF for %s",
            Path(path).name,
        )
        return _parse_pdf_pymupdf(path, cfg)
    return _parse_pdf_pymupdf(path, cfg)


def _docx_paragraphs_and_tables(path: str) -> tuple[list[str], list[dict], str]:
    doc = Document(path)
    paragraphs: list[str] = []
    tables: list[dict] = []
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paragraphs.append(t)
            parts.append(t)
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_rows.append({"row_num": row_idx + 1, "cells": cells})
                parts.append(" | ".join(cells))
        tables.append({"table_num": table_idx + 1, "rows": table_rows})
    return paragraphs, tables, "\n".join(parts).strip()


def _parse_docx_python_docx(path: str) -> dict:
    """запасной путь docx (python-docx), если docling недоступен или упал."""
    out = _make_base_result(path, "docx")
    paragraphs, tables, text = _docx_paragraphs_and_tables(path)
    out["text"] = _clean_whitespace(text)
    out["paragraphs"] = paragraphs
    out["tables"] = tables
    return out


def parse_docx(path: str) -> dict:
    if _docling_available():
        try:
            return _parse_with_docling(path, "docx", DEFAULT_PAGE_MARKER)
        except Exception as e:
            logger.warning(
                "Docling DOCX parse failed (%s), falling back to python-docx: %s",
                Path(path).name,
                e,
            )
            return _parse_docx_python_docx(path)
    logger.warning(
        "docling not installed; using python-docx for %s",
        Path(path).name,
    )
    return _parse_docx_python_docx(path)


def parse_txt(path: str) -> dict:
    out = _make_base_result(path, "txt")
    text = _read_text_with_fallbacks(path, ("utf-8", "cp1251", "latin-1")).strip()
    out["text"] = _clean_whitespace(text)
    return out


def parse_xlsx(path: str) -> dict:
    ext = Path(path).suffix.lower().lstrip(".") or "xlsx"
    out = _make_base_result(path, ext)
    excel_file = pd.ExcelFile(path)
    sheets: list[dict] = []
    full_text: list[str] = []

    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet_name)
        df = df.fillna("")
        columns = [str(c) for c in df.columns]
        rows_raw = df.astype(str).values.tolist()
        rows = [
            row
            for row in rows_raw
            if any(str(cell).strip() for cell in row)
        ]

        sheets.append({
            "sheet_name": sheet_name,
            "columns": columns,
            "rows": rows,
        })

        lines = [f"[Sheet: {sheet_name}]"]
        if columns:
            lines.append(" | ".join(columns))
        for row in rows:
            lines.append(" | ".join(str(cell).strip() for cell in row))
        full_text.append("\n".join(lines))

    out["text"] = _clean_whitespace("\n\n".join(full_text).strip())
    out["tables"] = sheets
    out["metadata"]["sheet_count"] = len(sheets)
    return out


def parse_rtf(path: str) -> dict:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as e:
        raise ValueError("RTF: установите пакет striprtf (pip install striprtf)") from e

    out = _make_base_result(path, "rtf")
    plain = _read_text_with_fallbacks(path, ("utf-8", "cp1251", "latin-1"))
    text = rtf_to_text(plain)
    out["text"] = _clean_whitespace(text)
    return out


def parse_doc(path: str) -> dict:
    out = _make_base_result(path, "doc")
    if platform.system() != "Darwin":
        raise ValueError(
            "Формат .doc на этой ОС не поддержан без внешнего конвертера. "
            "Сконвертируйте в .docx или используйте macOS (textutil)."
        )
    r = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", path],
        capture_output=True,
        text=True,
    )
    if r.returncode != 0:
        raise ValueError(f"textutil failed: {r.stderr or r.stdout}")
    out["text"] = _clean_whitespace(r.stdout)
    out["metadata"]["converter"] = "textutil"
    return out


def _ensure_ocr_input_path(path: str, ext: str) -> tuple[str, str | None]:
    if ext != ".heic":
        return path, None
    fd, tmp_path = tempfile.mkstemp(suffix=".png")
    import os

    os.close(fd)
    r = subprocess.run(["sips", "-s", "format", "png", path, "--out", tmp_path], capture_output=True, text=True)
    if r.returncode != 0:
        Path(tmp_path).unlink(missing_ok=True)
        raise ValueError(f"HEIC conversion failed: {r.stderr or r.stdout}")
    return tmp_path, tmp_path


def parse_image(path: str) -> dict:
    from ocr import get_ocr, run_ocr_on_image

    ext = Path(path).suffix.lower()
    file_type = ext.lstrip(".")
    out = _make_base_result(path, file_type)
    ocr = get_ocr()
    ocr_path, temp_path = _ensure_ocr_input_path(path, ext)
    try:
        result = run_ocr_on_image(ocr, ocr_path)
    finally:
        if temp_path:
            Path(temp_path).unlink(missing_ok=True)

    lines = []
    for block in result or []:
        if not block:
            continue
        for item in block:
            t = item[1][0].strip()
            if t:
                lines.append(t)

    out["text"] = _clean_whitespace(postprocess_ocr_text("\n".join(lines)))
    out["metadata"]["ocr_used"] = True
    out["metadata"]["ocr_engine"] = "paddleocr"
    out["metadata"]["ocr_pages"] = 1
    return out


def _extract_for_extension(
    path: str,
    ext: str,
    pdf_config: PdfExtractionConfig | None = None,
) -> dict:
    """маршрут по расширению: pdf/docx — внутри своих парсеров (docling при настройке), таблицы — pandas, картинки — ocr."""
    if ext == ".pdf":
        return parse_pdf(path, pdf_config)
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".doc":
        return parse_doc(path)
    if ext == ".rtf":
        return parse_rtf(path)
    if ext == ".txt":
        return parse_txt(path)
    if ext in {".xls", ".xlsx"}:
        return parse_xlsx(path)
    if ext in IMAGE_EXTENSIONS:
        return parse_image(path)
    raise ValueError(f"Unsupported file type: {ext}")


def extract_text(path: str, pdf_config: PdfExtractionConfig | None = None) -> dict:
    ext = Path(path).suffix.lower()
    return _extract_for_extension(path, ext, pdf_config)
